#!/usr/bin/env python3
"""
Markdown to DOCX Converter - Fixed Image Display Issues

This script provides multiple methods to convert Markdown files to DOCX format.
主要修复了图片显示问题。
"""

import os
import sys
import argparse
from pathlib import Path
from typing import Optional
import re
import tempfile
import shutil


def validate_image_file(img_path: str) -> bool:
    """验证图片文件是否有效"""
    if not os.path.exists(img_path):
        return False
    
    # 检查文件是否为空
    if os.path.getsize(img_path) == 0:
        return False
    
    # 检查文件扩展名
    valid_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'}
    ext = os.path.splitext(img_path)[1].lower()
    return ext in valid_extensions


def convert_image_format(img_path: str) -> str:
    """转换图片格式为DOCX兼容的格式"""
    try:
        from PIL import Image
        
        # 支持的格式
        supported_formats = {'.jpg', '.jpeg', '.png'}
        ext = os.path.splitext(img_path)[1].lower()
        
        if ext in supported_formats:
            return img_path
        
        # 转换为PNG格式
        with Image.open(img_path) as img:
            # 如果是RGBA模式但要保存为JPEG，需要转换
            if img.mode == 'RGBA' and ext.lower() in ['.jpg', '.jpeg']:
                # 创建白色背景
                background = Image.new('RGB', img.size, (255, 255, 255))
                background.paste(img, mask=img.split()[-1])  # 使用alpha通道作为mask
                img = background
            
            converted_path = os.path.splitext(img_path)[0] + '_converted.png'
            img.save(converted_path, 'PNG')
            return converted_path
            
    except ImportError:
        print("⚠️ PIL/Pillow not installed. Some image formats may not work.")
        print("Install with: pip install Pillow")
        return img_path
    except Exception as e:
        print(f"⚠️ Error converting image format: {e}")
        return img_path


def method_2_markdown_docx_fixed(md_file: str, output_file: str) -> bool:
    """
    改进的Method 2: 修复图片显示问题
    """
    try:
        import markdown
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import qn
        import requests
        from urllib.parse import urlparse, urljoin
        
        # Read markdown file
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Create new document
        doc = Document()
        
        # Set default font functions
        def set_font_style(run, font_name="宋体", font_size=12):
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        
        def set_paragraph_font(paragraph, font_name="宋体", font_size=12):
            for run in paragraph.runs:
                set_font_style(run, font_name, font_size)
        
        # 改进的图片处理函数
        def process_image_enhanced(line, base_path):
            # Extract image info from markdown: ![alt](src)
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if not match:
                return False
            
            alt_text = match.group(1)
            img_src = match.group(2).strip()
            
            # Remove angle brackets if present
            if img_src.startswith('<') and img_src.endswith('>'):
                img_src = img_src[1:-1]
            
            # Remove quotes if present
            img_src = img_src.strip('\'"')
            
            # 处理带参数的图片URL/路径（如CDN参数）
            # 例如: image.jpg!t340x226 或 image.jpg?width=340&height=226
            original_img_src = img_src
            clean_img_src = img_src
            
            # 检查是否包含CDN参数（感叹号或问号）
            if '!' in img_src:
                clean_img_src = img_src.split('!')[0]
                print(f"🔧 Detected CDN parameters, clean path: {clean_img_src}")
            elif '?' in img_src and not img_src.startswith(('http://', 'https://')):
                clean_img_src = img_src.split('?')[0]
                print(f"🔧 Detected URL parameters, clean path: {clean_img_src}")
            
            print(f"🖼️ Processing image: {alt_text} -> {original_img_src}")
            if clean_img_src != original_img_src:
                print(f"    Clean path: {clean_img_src}")
            
            temp_files = []  # 记录临时文件，用于清理
            
            try:
                img_path = None
                
                # Check if it's a URL
                if original_img_src.startswith(('http://', 'https://')):
                    print(f"🌐 Downloading image from URL: {original_img_src}")
                    try:
                        headers = {
                            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                        }
                        response = requests.get(original_img_src, timeout=30, headers=headers, stream=True)
                        response.raise_for_status()
                        
                        # 确定文件扩展名
                        content_type = response.headers.get('content-type', '').lower()
                        if 'jpeg' in content_type or 'jpg' in content_type:
                            ext = '.jpg'
                        elif 'png' in content_type:
                            ext = '.png'
                        elif 'gif' in content_type:
                            ext = '.gif'
                        elif 'webp' in content_type:
                            ext = '.webp'
                        else:
                            # 尝试从URL获取扩展名（使用clean_img_src去掉参数）
                            parsed_url = urlparse(clean_img_src)
                            url_ext = os.path.splitext(parsed_url.path)[1].lower()
                            ext = url_ext if url_ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp'] else '.jpg'
                        
                        # 创建临时文件
                        temp_fd, temp_path = tempfile.mkstemp(suffix=ext)
                        temp_files.append(temp_path)
                        
                        # 写入文件
                        with os.fdopen(temp_fd, 'wb') as f:
                            for chunk in response.iter_content(chunk_size=8192):
                                f.write(chunk)
                        
                        img_path = temp_path
                        print(f"✅ Downloaded to: {temp_path}")
                        
                    except Exception as e:
                        print(f"❌ Could not download image {original_img_src}: {e}")
                        return False
                
                else:
                    # Local file path - 使用clean_img_src（去掉参数后的路径）
                    possible_paths = []
                    
                    # 绝对路径
                    if os.path.isabs(clean_img_src):
                        possible_paths.append(clean_img_src)
                    else:
                        # 相对于markdown文件的路径
                        possible_paths.append(os.path.join(base_path, clean_img_src))
                        
                        # 相对于当前工作目录的路径
                        possible_paths.append(os.path.join(os.getcwd(), clean_img_src))
                        
                        # 处理可能的路径分隔符问题
                        img_src_normalized = clean_img_src.replace('\\', os.sep).replace('/', os.sep)
                        if img_src_normalized != clean_img_src:
                            possible_paths.append(os.path.join(base_path, img_src_normalized))
                            possible_paths.append(os.path.join(os.getcwd(), img_src_normalized))
                    
                    # 去除重复路径
                    possible_paths = list(dict.fromkeys(possible_paths))  # 保持顺序的去重
                    
                    print(f"🔍 Searching for image in paths:")
                    for path in possible_paths:
                        print(f"    {path}")
                    
                    for path in possible_paths:
                        if os.path.exists(path) and validate_image_file(path):
                            img_path = path
                            print(f"✅ Found valid image at: {img_path}")
                            break
                    
                    if not img_path:
                        print(f"❌ Image file not found or invalid in any searched paths")
                        return False
                
                # 验证和转换图片格式
                if not validate_image_file(img_path):
                    print(f"❌ Invalid image file: {img_path}")
                    return False
                
                # 尝试转换格式以确保兼容性
                converted_path = convert_image_format(img_path)
                if converted_path != img_path:
                    temp_files.append(converted_path)
                    img_path = converted_path
                
                # 插入图片到文档
                try:
                    # 获取图片信息
                    from PIL import Image
                    with Image.open(img_path) as img:
                        width, height = img.size
                        print(f"📏 Image size: {width}x{height}")
                    
                    # 计算合适的显示尺寸（最大6英寸宽度）
                    max_width = Inches(6)
                    if width > height:
                        display_width = min(max_width, Inches(width / 100))  # 假设100px = 1inch
                    else:
                        display_width = min(max_width, Inches(4))
                    
                except:
                    # 如果无法获取图片尺寸，使用默认值
                    display_width = Inches(4)
                
                # 添加图片段落
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
                
                run = paragraph.add_run()
                run.add_picture(img_path, width=display_width)
                
                # 添加图片说明（如果有alt text）
                if alt_text.strip():
                    caption_paragraph = doc.add_paragraph(f"图片: {alt_text}")
                    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_paragraph_font(caption_paragraph, "宋体", 10)
                
                print(f"✅ Image inserted successfully")
                return True
                
            except Exception as e:
                print(f"❌ Error processing image {original_img_src}: {e}")
                import traceback
                traceback.print_exc()
                return False
            
            finally:
                # 清理临时文件
                for temp_file in temp_files:
                    try:
                        if os.path.exists(temp_file):
                            os.remove(temp_file)
                            print(f"🧹 Cleaned up temp file: {temp_file}")
                    except Exception as e:
                        print(f"⚠️ Could not clean up temp file {temp_file}: {e}")
        
        # Get base path for relative image paths
        base_path = os.path.dirname(os.path.abspath(md_file))
        print(f"📁 Base path for images: {base_path}")
        
        # Split content by lines and process
        lines = md_content.split('\n')
        in_code_block = False
        code_content = []
        
        for line in lines:
            original_line = line
            line_stripped = line.strip()
            
            # Handle code blocks
            if line_stripped.startswith('```'):
                if in_code_block:
                    # End of code block
                    if code_content:
                        code_text = '\n'.join(code_content)
                        p = doc.add_paragraph(code_text)
                        # 设置代码样式
                        for run in p.runs:
                            run.font.name = 'Consolas'
                            run.font.size = Pt(10)
                    code_content = []
                    in_code_block = False
                else:
                    # Start of code block
                    in_code_block = True
                continue
            
            if in_code_block:
                code_content.append(original_line)
                continue
            
            if not line_stripped:
                # 空行也添加，保持文档结构
                doc.add_paragraph()
                continue
                
            # Headers with different font sizes
            if line_stripped.startswith('# '):
                heading = doc.add_heading(line_stripped[2:], level=1)
                set_paragraph_font(heading, "宋体", 18)
            elif line_stripped.startswith('## '):
                heading = doc.add_heading(line_stripped[3:], level=2)
                set_paragraph_font(heading, "宋体", 16)
            elif line_stripped.startswith('### '):
                heading = doc.add_heading(line_stripped[4:], level=3)
                set_paragraph_font(heading, "宋体", 14)
            elif line_stripped.startswith('#### '):
                heading = doc.add_heading(line_stripped[5:], level=4)
                set_paragraph_font(heading, "宋体", 13)
            elif line_stripped.startswith('##### '):
                heading = doc.add_heading(line_stripped[6:], level=5)
                set_paragraph_font(heading, "宋体", 12)
            elif line_stripped.startswith('###### '):
                heading = doc.add_heading(line_stripped[7:], level=6)
                set_paragraph_font(heading, "宋体", 12)
            # Lists
            elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
                p = doc.add_paragraph(line_stripped[2:], style='List Bullet')
                set_paragraph_font(p, "宋体", 12)
            elif re.match(r'^\d+\. ', line_stripped):
                content = line_stripped[line_stripped.index(' ')+1:]
                p = doc.add_paragraph(content, style='List Number')
                set_paragraph_font(p, "宋体", 12)
            # Images - 使用改进的处理函数
            elif line_stripped.startswith('!['):
                success = process_image_enhanced(line_stripped, base_path)
                if not success:
                    # 如果图片处理失败，添加占位文本
                    match = re.match(r'!\[(.*?)\]\((.*?)\)', line_stripped)
                    if match:
                        alt_text = match.group(1)
                        img_src = match.group(2)
                        p = doc.add_paragraph(f"[图片未找到: {alt_text} - 路径: {img_src}]")
                        set_paragraph_font(p, "宋体", 12)
                        # 设置为红色以突出显示
                        try:
                            for run in p.runs:
                                run.font.color.rgb = RGBColor(255, 0, 0)
                        except Exception:
                            pass  # 如果设置颜色失败，忽略
            # Regular paragraphs
            else:
                p = doc.add_paragraph(line_stripped)
                set_paragraph_font(p, "宋体", 12)
        
        # Save document
        doc.save(output_file)
        print(f"✅ Successfully converted using enhanced markdown+docx: {output_file}")
        return True
        
    except ImportError as e:
        print(f"❌ Required libraries not installed: {e}")
        print("Install with: pip install markdown python-docx requests Pillow")
        return False
    except Exception as e:
        print(f"❌ Error with enhanced markdown+docx method: {e}")
        import traceback
        traceback.print_exc()
        return False


def convert_md_to_docx_fixed(md_file: str, output_file: Optional[str] = None) -> bool:
    """
    使用修复版本的转换函数
    """
    # Validate input file
    if not os.path.exists(md_file):
        print(f"❌ Input file not found: {md_file}")
        return False
    
    # Generate output filename if not provided
    if output_file is None:
        output_file = Path(md_file).with_suffix('.docx')
    
    print(f"🔄 Converting: {md_file} -> {output_file}")
    print("📋 This version includes enhanced image processing:")
    print("  • Better path resolution")
    print("  • Image format validation and conversion")
    print("  • Improved error handling")
    print("  • Temporary file cleanup")
    
    return method_2_markdown_docx_fixed(md_file, str(output_file))


def install_dependencies_fixed():
    """Install required dependencies including PIL for image processing"""
    import subprocess
    
    packages = [
        'markdown',
        'python-docx',
        'requests',
        'Pillow',  # 添加PIL用于图片处理
    ]
    
    print("📦 Installing dependencies...")
    for package in packages:
        try:
            subprocess.check_call([sys.executable, '-m', 'pip', 'install', package])
            print(f"✅ Installed: {package}")
        except subprocess.CalledProcessError:
            print(f"❌ Failed to install: {package}")


def main():
    """Main function with command line interface"""
    parser = argparse.ArgumentParser(description='Convert Markdown files to DOCX format (Fixed Image Issues)')
    parser.add_argument('input', nargs='?', help='Input markdown file')
    parser.add_argument('-o', '--output', help='Output DOCX file (optional)')
    parser.add_argument('--install-deps', action='store_true',
                       help='Install required dependencies')
    parser.add_argument('--debug', action='store_true',
                       help='Enable debug output')
    
    args = parser.parse_args()
    
    if args.install_deps:
        install_dependencies_fixed()
        return
    
    if not args.input:
        print("🚀 Markdown to DOCX Converter (Image Issues Fixed)")
        print("=" * 55)
        print("\n🔧 Fixed Issues:")
        print("  • 改进图片路径解析")
        print("  • 添加图片格式验证和转换")
        print("  • 增强错误处理和调试信息")
        print("  • 自动清理临时文件")
        print("  • 支持更多图片格式")
        print("\n✨ Features:")
        print("  • 图片自动插入（支持本地文件和URL）")
        print("  • 统一宋体字体设置")
        print("  • 标题和正文字体大小自动区分")
        print("  • 图片居中对齐和尺寸自适应")
        print("\n📏 字体设置:")
        print("  • 标题: 宋体 H1(18pt) H2(16pt) H3(14pt) H4(13pt) H5/H6(12pt)")
        print("  • 正文: 宋体 12pt")
        print("  • 代码: Consolas 10pt")
        print("\nUsage examples:")
        print("  python md_to_docx_fixed.py document.md")
        print("  python md_to_docx_fixed.py document.md -o output.docx")
        print("  python md_to_docx_fixed.py --install-deps")
        
        # Interactive mode
        md_file = input("\n📄 Enter markdown file path: ").strip()
        if md_file and os.path.exists(md_file):
            success = convert_md_to_docx_fixed(md_file)
            if success:
                print("✅ Conversion completed successfully!")
            else:
                print("❌ Conversion failed. Check the output above for details.")
        elif md_file:
            print(f"❌ File not found: {md_file}")
        return
    
    success = convert_md_to_docx_fixed(args.input, args.output)
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()