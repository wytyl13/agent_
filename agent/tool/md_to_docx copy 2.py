#!/usr/bin/env python3
"""
Markdown to DOCX Converter - Custom Format Version

This script provides multiple methods to convert Markdown files to DOCX format.
自定义格式版本：蓝色标题、自动编号、指定字体大小
"""

import os
import sys
import argparse
from pathlib import Path
from typing import Optional
import re
import tempfile
import shutil


def validate_image_file(img_path: str) -> bool:
    """验证图片文件是否有效"""
    if not os.path.exists(img_path):
        return False
    
    # 检查文件是否为空
    if os.path.getsize(img_path) == 0:
        return False
    
    # 检查文件扩展名
    valid_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'}
    ext = os.path.splitext(img_path)[1].lower()
    return ext in valid_extensions


def convert_image_format(img_path: str) -> str:
    """转换图片格式为DOCX兼容的格式"""
    try:
        from PIL import Image
        
        # 支持的格式
        supported_formats = {'.jpg', '.jpeg', '.png'}
        ext = os.path.splitext(img_path)[1].lower()
        
        if ext in supported_formats:
            return img_path
        
        # 转换为PNG格式
        with Image.open(img_path) as img:
            # 如果是RGBA模式但要保存为JPEG，需要转换
            if img.mode == 'RGBA' and ext.lower() in ['.jpg', '.jpeg']:
                # 创建白色背景
                background = Image.new('RGB', img.size, (255, 255, 255))
                background.paste(img, mask=img.split()[-1])  # 使用alpha通道作为mask
                img = background
            
            converted_path = os.path.splitext(img_path)[0] + '_converted.png'
            img.save(converted_path, 'PNG')
            return converted_path
            
    except ImportError:
        print("⚠️ PIL/Pillow not installed. Some image formats may not work.")
        print("Install with: pip install Pillow")
        return img_path
    except Exception as e:
        print(f"⚠️ Error converting image format: {e}")
        return img_path


def method_2_markdown_docx_custom(md_file: str, output_file: str) -> bool:
    """
    自定义格式版本: 蓝色标题、自动编号、指定字体大小
    """
    try:
        import markdown
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import qn
        import requests
        from urllib.parse import urlparse, urljoin
        
        # Read markdown file
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Create new document
        doc = Document()
        
        # 定义蓝色颜色 (根据图片中的蓝色，大概是这个颜色值)
        BLUE_COLOR = RGBColor(0, 112, 192)  # 深蓝色
        
        # Set default font functions
        def set_font_style(run, font_name="宋体", font_size=15, color=None):
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            if color:
                run.font.color.rgb = color
        
        def set_paragraph_font(paragraph, font_name="宋体", font_size=15, color=None):
            for run in paragraph.runs:
                set_font_style(run, font_name, font_size, color)
        
        # 图片URL/路径清理和验证函数
        def clean_and_validate_image_url(img_src):
            """清理图片URL并验证格式"""
            original_img_src = img_src.strip()
            
            # Remove angle brackets if present - 但保留用于后续处理
            has_angle_brackets = False
            if original_img_src.startswith('<') and original_img_src.endswith('>'):
                has_angle_brackets = True
                original_img_src = original_img_src[1:-1]
            
            # Remove quotes if present
            original_img_src = original_img_src.strip('\'"')
            
            # 支持的图片格式
            valid_image_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'}
            
            # 检查是否为URL
            is_url = original_img_src.startswith(('http://', 'https://'))
            
            clean_img_src = original_img_src
            
            if is_url:
                # 处理网络URL
                from urllib.parse import urlparse, urlunparse
                parsed = urlparse(original_img_src)
                
                # 获取路径部分的扩展名
                path_ext = os.path.splitext(parsed.path)[1].lower()
                
                # 如果路径有有效的图片扩展名，清理参数
                if path_ext in valid_image_extensions:
                    # 清理URL参数
                    clean_parsed = parsed._replace(query='', fragment='')
                    clean_img_src = urlunparse(clean_parsed)
                    if clean_img_src != original_img_src:
                        print(f"🔧 Cleaned URL parameters: {original_img_src} -> {clean_img_src}")
                else:
                    print(f"⚠️ URL does not have valid image extension: {original_img_src}")
            else:
                # 处理本地文件路径
                print(f"🔍 Original path: {original_img_src}")
                
                # 先处理各种参数格式
                if '!m>' in original_img_src:
                    # 处理 !m> 这种特殊格式
                    clean_img_src = original_img_src.split('!m>')[0]
                    print(f"🔧 Detected !m> parameters, clean path: {clean_img_src}")
                elif '!' in original_img_src:
                    # 处理其他感叹号参数
                    clean_img_src = original_img_src.split('!')[0]
                    print(f"🔧 Detected CDN parameters, clean path: {clean_img_src}")
                elif '?' in original_img_src:
                    # 处理问号参数
                    clean_img_src = original_img_src.split('?')[0]
                    print(f"🔧 Detected URL parameters, clean path: {clean_img_src}")
                
                # 验证文件扩展名
                path_ext = os.path.splitext(clean_img_src)[1].lower()
                if path_ext not in valid_image_extensions:
                    print(f"⚠️ Path does not have valid image extension: {clean_img_src}")
                    # 尝试添加常见的图片扩展名
                    for ext in ['.jpg', '.png', '.jpeg']:
                        test_path = clean_img_src + ext
                        print(f"🔍 Testing path with extension: {test_path}")
                        # 检查多个可能的位置
                        test_locations = [
                            test_path,
                            os.path.join(base_path, test_path) if not os.path.isabs(test_path) else test_path,
                            os.path.join(os.getcwd(), test_path) if not os.path.isabs(test_path) else test_path
                        ]
                        for location in test_locations:
                            if os.path.exists(location):
                                print(f"🔧 Found image with extension: {location}")
                                clean_img_src = test_path
                                break
                        if clean_img_src == test_path:
                            break
            
            # 如果原来有角括号，在返回的路径中也加上（用于后续识别）
            return_original = f"<{original_img_src}>" if has_angle_brackets else original_img_src
            return_clean = f"<{clean_img_src}>" if has_angle_brackets else clean_img_src
            
            return return_original, return_clean
        
        # 改进的图片处理函数
        def process_image_enhanced(line, base_path):
            # Extract image info from markdown: ![alt](src)
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if not match:
                return False
            
            alt_text = match.group(1)
            img_src = match.group(2).strip()
            
            # 清理和验证图片URL/路径
            original_img_src, clean_img_src = clean_and_validate_image_url(img_src)
            
            print(f"🖼️ Processing image: {alt_text} -> {original_img_src}")
            if clean_img_src != original_img_src:
                print(f"    Clean path: {clean_img_src}")
        # 改进的图片处理函数
        def process_image_enhanced(line, base_path):
            # Extract image info from markdown: ![alt](src)
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if not match:
                return False
            
            alt_text = match.group(1)
            img_src = match.group(2).strip()
            
            print(f"🖼️ Processing image: {alt_text}")
            print(f"📍 Raw image source: {img_src}")
            
            # 清理和验证图片URL/路径
            original_img_src, clean_img_src = clean_and_validate_image_url(img_src)
            
            # 去掉角括号进行实际文件操作
            actual_original = original_img_src.strip('<>')
            actual_clean = clean_img_src.strip('<>')
            
            print(f"🔍 Actual original path: {actual_original}")
            print(f"🔍 Actual clean path: {actual_clean}")
            
            temp_files = []  # 记录临时文件，用于清理
            
            try:
                img_path = None
                
                # Check if it's a URL
                if actual_clean.startswith(('http://', 'https://')):
                    print(f"🌐 Downloading image from URL: {actual_clean}")
                    try:
                        headers = {
                            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                        }
                        response = requests.get(actual_clean, timeout=30, headers=headers, stream=True)
                        response.raise_for_status()
                        
                        # 确定文件扩展名
                        content_type = response.headers.get('content-type', '').lower()
                        if 'jpeg' in content_type or 'jpg' in content_type:
                            ext = '.jpg'
                        elif 'png' in content_type:
                            ext = '.png'
                        elif 'gif' in content_type:
                            ext = '.gif'
                        elif 'webp' in content_type:
                            ext = '.webp'
                        else:
                            # 尝试从URL获取扩展名（使用actual_clean去掉参数）
                            parsed_url = urlparse(actual_clean)
                            url_ext = os.path.splitext(parsed_url.path)[1].lower()
                            ext = url_ext if url_ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp'] else '.jpg'
                        
                        # 创建临时文件
                        temp_fd, temp_path = tempfile.mkstemp(suffix=ext)
                        temp_files.append(temp_path)
                        
                        # 写入文件
                        with os.fdopen(temp_fd, 'wb') as f:
                            for chunk in response.iter_content(chunk_size=8192):
                                f.write(chunk)
                        
                        img_path = temp_path
                        print(f"✅ Downloaded to: {temp_path}")
                        
                    except Exception as e:
                        print(f"❌ Could not download image {actual_clean}: {e}")
                        # 如果清理后的URL失败，尝试原始URL
                        if actual_clean != actual_original:
                            print(f"🔄 Trying original URL: {actual_original}")
                            try:
                                response = requests.get(actual_original, timeout=30, headers=headers, stream=True)
                                response.raise_for_status()
                                
                                # 使用原始URL下载逻辑...
                                content_type = response.headers.get('content-type', '').lower()
                                if 'jpeg' in content_type or 'jpg' in content_type:
                                    ext = '.jpg'
                                elif 'png' in content_type:
                                    ext = '.png'
                                elif 'gif' in content_type:
                                    ext = '.gif'
                                elif 'webp' in content_type:
                                    ext = '.webp'
                                else:
                                    ext = '.jpg'
                                
                                temp_fd, temp_path = tempfile.mkstemp(suffix=ext)
                                temp_files.append(temp_path)
                                
                                with os.fdopen(temp_fd, 'wb') as f:
                                    for chunk in response.iter_content(chunk_size=8192):
                                        f.write(chunk)
                                
                                img_path = temp_path
                                print(f"✅ Downloaded using original URL to: {temp_path}")
                            except Exception as e2:
                                print(f"❌ Both clean and original URLs failed: {e2}")
                                return False
                
                else:
                    # Local file path - 使用actual_clean（去掉参数后的路径）
                    possible_paths = []
                    
                    # 绝对路径
                    if os.path.isabs(actual_clean):
                        possible_paths.append(actual_clean)
                    else:
                        # 相对于markdown文件的路径
                        possible_paths.append(os.path.join(base_path, actual_clean))
                        
                        # 相对于当前工作目录的路径
                        possible_paths.append(os.path.join(os.getcwd(), actual_clean))
                        
                        # 处理可能的路径分隔符问题
                        img_src_normalized = actual_clean.replace('\\', os.sep).replace('/', os.sep)
                        if img_src_normalized != actual_clean:
                            possible_paths.append(os.path.join(base_path, img_src_normalized))
                            possible_paths.append(os.path.join(os.getcwd(), img_src_normalized))
                    
                    # 如果清理后的路径与原始路径不同，也尝试原始路径
                    if actual_clean != actual_original:
                        if os.path.isabs(actual_original):
                            possible_paths.append(actual_original)
                        else:
                            possible_paths.append(os.path.join(base_path, actual_original))
                            possible_paths.append(os.path.join(os.getcwd(), actual_original))
                    
                    # 去除重复路径
                    possible_paths = list(dict.fromkeys(possible_paths))  # 保持顺序的去重
                    
                    print(f"🔍 Searching for clean image in paths:")
                    for path in possible_paths:
                        print(f"    {path}")
                    
                    # 首先尝试清理后的路径
                    for path in possible_paths:
                        if os.path.exists(path) and validate_image_file(path):
                            img_path = path
                            print(f"✅ Found valid image at: {img_path}")
                            break
                    
                    # 如果没找到清理后的路径，尝试找原始路径并重命名
                    if not img_path and actual_clean != actual_original:
                        print(f"🔍 Clean path not found, searching for original path to rename...")
                        print(f"🔍 Looking for original file: {actual_original}")
                        
                        # 构建原始路径的可能位置
                        original_possible_paths = []
                        if os.path.isabs(actual_original):
                            original_possible_paths.append(actual_original)
                        else:
                            original_possible_paths.append(os.path.join(base_path, actual_original))
                            original_possible_paths.append(os.path.join(os.getcwd(), actual_original))
                        
                        # 去除重复路径
                        original_possible_paths = list(dict.fromkeys(original_possible_paths))
                        
                        print(f"🔍 Searching for original image in paths:")
                        for path in original_possible_paths:
                            print(f"    {path}")
                        
                        for original_path in original_possible_paths:
                            if os.path.exists(original_path) and validate_image_file(original_path):
                                print(f"✅ Found original image at: {original_path}")
                                
                                # 计算清理后的路径
                                if os.path.isabs(actual_clean):
                                    target_clean_path = actual_clean
                                else:
                                    # 使用与原始文件相同的目录
                                    target_clean_path = os.path.join(os.path.dirname(original_path), os.path.basename(actual_clean))
                                
                                print(f"🎯 Target clean path: {target_clean_path}")
                                
                                try:
                                    # 重命名文件
                                    if not os.path.exists(target_clean_path):
                                        shutil.move(original_path, target_clean_path)
                                        print(f"📝 Renamed image: {original_path} -> {target_clean_path}")
                                        img_path = target_clean_path
                                    else:
                                        print(f"⚠️ Clean path already exists: {target_clean_path}")
                                        print(f"🔄 Using existing clean file")
                                        img_path = target_clean_path
                                    break
                                except Exception as e:
                                    print(f"❌ Could not rename file: {e}")
                                    print(f"🔄 Trying to copy instead...")
                                    try:
                                        if not os.path.exists(target_clean_path):
                                            shutil.copy2(original_path, target_clean_path)
                                            print(f"📋 Copied image: {original_path} -> {target_clean_path}")
                                            img_path = target_clean_path
                                        else:
                                            img_path = target_clean_path
                                        break
                                    except Exception as e2:
                                        print(f"❌ Could not copy file either: {e2}")
                                        # 使用原始文件
                                        print(f"🔄 Using original file: {original_path}")
                                        img_path = original_path
                                        break
                    
                    if not img_path:
                        print(f"❌ Image file not found or invalid in any searched paths")
                        return False
                
                # 验证和转换图片格式
                if not validate_image_file(img_path):
                    print(f"❌ Invalid image file: {img_path}")
                    return False
                
                # 尝试转换格式以确保兼容性
                converted_path = convert_image_format(img_path)
                if converted_path != img_path:
                    temp_files.append(converted_path)
                    img_path = converted_path
                
                # 插入图片到文档
                try:
                    # 获取图片信息
                    from PIL import Image
                    with Image.open(img_path) as img:
                        width, height = img.size
                        print(f"📏 Image size: {width}x{height}")
                    
                    # 计算合适的显示尺寸（最大6英寸宽度）
                    max_width = Inches(6)
                    if width > height:
                        display_width = min(max_width, Inches(width / 100))  # 假设100px = 1inch
                    else:
                        display_width = min(max_width, Inches(4))
                    
                except:
                    # 如果无法获取图片尺寸，使用默认值
                    display_width = Inches(4)
                
                # 添加图片段落
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
                
                run = paragraph.add_run()
                run.add_picture(img_path, width=display_width)
                
                # 不添加图片说明文字
                
                print(f"✅ Image inserted successfully")
                return True
                
            except Exception as e:
                print(f"❌ Error processing image {original_img_src}: {e}")
                import traceback
                traceback.print_exc()
                return False
            
            finally:
                # 清理临时文件
                for temp_file in temp_files:
                    try:
                        if os.path.exists(temp_file):
                            os.remove(temp_file)
                            print(f"🧹 Cleaned up temp file: {temp_file}")
                    except Exception as e:
                        print(f"⚠️ Could not clean up temp file {temp_file}: {e}")
        
        # Get base path for relative image paths
        base_path = os.path.dirname(os.path.abspath(md_file))
        print(f"📁 Base path for images: {base_path}")
        
        # Split content by lines and process
        lines = md_content.split('\n')
        in_code_block = False
        code_content = []
        
        # 计数器用于三级标题编号 - 在每个二级标题下重新计数
        h3_counter = 0
        
        for line in lines:
            original_line = line
            line_stripped = line.strip()
            
            # Handle code blocks
            if line_stripped.startswith('```'):
                if in_code_block:
                    # End of code block
                    if code_content:
                        code_text = '\n'.join(code_content)
                        p = doc.add_paragraph(code_text)
                        # 设置代码样式
                        for run in p.runs:
                            run.font.name = 'Consolas'
                            run.font.size = Pt(10)
                    code_content = []
                    in_code_block = False
                else:
                    # Start of code block
                    in_code_block = True
                continue
            
            if in_code_block:
                code_content.append(original_line)
                continue
            
            if not line_stripped:
                # 空行也添加，保持文档结构
                doc.add_paragraph()
                continue
                
            # Headers with custom formatting
            if line_stripped.startswith('# '):
                # 一级标题：保持原样，不做特殊处理
                heading = doc.add_heading(line_stripped[2:], level=1)
                set_paragraph_font(heading, "宋体", 20, BLUE_COLOR)
            elif line_stripped.startswith('## '):
                # 二级标题：宋体20号字体，蓝色，无编号，重置三级标题计数器
                h3_counter = 0  # 重置三级标题计数器
                heading = doc.add_heading(line_stripped[3:], level=2)
                set_paragraph_font(heading, "宋体", 20, BLUE_COLOR)
            elif line_stripped.startswith('### '):
                # 三级标题：自动编号，宋体18号字体，蓝色
                h3_counter += 1
                heading_text = f"{h3_counter:02d} {line_stripped[4:]}"  # 01 02 03格式
                heading = doc.add_heading(heading_text, level=3)
                set_paragraph_font(heading, "宋体", 18, BLUE_COLOR)
            elif line_stripped.startswith('#### '):
                # 四级标题：宋体16号字体，蓝色
                heading = doc.add_heading(line_stripped[5:], level=4)
                set_paragraph_font(heading, "宋体", 16, BLUE_COLOR)
            elif line_stripped.startswith('##### '):
                # 五级标题：宋体15号字体，蓝色
                heading = doc.add_heading(line_stripped[6:], level=5)
                set_paragraph_font(heading, "宋体", 15, BLUE_COLOR)
            elif line_stripped.startswith('###### '):
                # 六级标题：宋体15号字体，蓝色
                heading = doc.add_heading(line_stripped[7:], level=6)
                set_paragraph_font(heading, "宋体", 15, BLUE_COLOR)
            # Lists
            elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
                p = doc.add_paragraph(line_stripped[2:], style='List Bullet')
                set_paragraph_font(p, "宋体", 15)  # 正文15号字体
            elif re.match(r'^\d+\. ', line_stripped):
                content = line_stripped[line_stripped.index(' ')+1:]
                p = doc.add_paragraph(content, style='List Number')
                set_paragraph_font(p, "宋体", 15)  # 正文15号字体
            # Images - 使用改进的处理函数
            elif line_stripped.startswith('!['):
                success = process_image_enhanced(line_stripped, base_path)
                if not success:
                    # 如果图片处理失败，添加占位文本
                    match = re.match(r'!\[(.*?)\]\((.*?)\)', line_stripped)
                    if match:
                        alt_text = match.group(1)
                        img_src = match.group(2)
                        _, clean_src = clean_and_validate_image_url(img_src)
                        p = doc.add_paragraph(f"[图片未找到: {alt_text} - 原始路径: {img_src} - 清理后路径: {clean_src}]")
                        set_paragraph_font(p, "宋体", 15)
                        # 设置为红色以突出显示
                        try:
                            for run in p.runs:
                                run.font.color.rgb = RGBColor(255, 0, 0)
                        except Exception:
                            pass  # 如果设置颜色失败，忽略
            # Regular paragraphs - 正文宋体15号字体，跳过横线分割符
            else:
                # 跳过横线分割符（如 ---、___、***）
                if line_stripped in ['---', '___', '***'] or re.match(r'^[-_*]{3,}$', line_stripped):
                    continue
                
                p = doc.add_paragraph(line_stripped)
                set_paragraph_font(p, "宋体", 15)  # 正文15号字体
        
        # Save document
        doc.save(output_file)
        print(f"✅ Successfully converted using custom format: {output_file}")
        print(f"📝 Applied custom formatting:")
        print(f"   • 一级标题: 保持原样")
        print(f"   • 二级标题: 宋体 20pt 蓝色 (无编号)")
        print(f"   • 三级标题: 宋体 18pt 蓝色 + 自动编号(每个二级标题下从01开始)")
        print(f"   • 四级及以下标题: 宋体 15-16pt 蓝色")
        print(f"   • 正文: 宋体 15pt 黑色")
        print(f"   • 图片: 无说明文字，居中显示，智能URL清理")
        print(f"   • 分割线: 自动忽略")
        return True
        
    except ImportError as e:
        print(f"❌ Required libraries not installed: {e}")
        print("Install with: pip install markdown python-docx requests Pillow")
        return False
    except Exception as e:
        print(f"❌ Error with custom format method: {e}")
        import traceback
        traceback.print_exc()
        return False


def convert_md_to_docx_custom(md_file: str, output_file: Optional[str] = None) -> bool:
    """
    使用自定义格式版本的转换函数
    """
    # Validate input file
    if not os.path.exists(md_file):
        print(f"❌ Input file not found: {md_file}")
        return False
    
    # Generate output filename if not provided
    if output_file is None:
        output_file = Path(md_file).with_suffix('.docx')
    
    print(f"🔄 Converting: {md_file} -> {output_file}")
    print("📋 Custom Format Features:")
    print("  • 一级标题: 保持原样")
    print("  • 二级标题: 宋体 20pt 蓝色 (无编号)")
    print("  • 三级标题: 宋体 18pt 蓝色 + 自动编号 (每个二级标题下从01开始)")
    print("  • 四级标题: 宋体 16pt 蓝色")
    print("  • 正文内容: 宋体 15pt 黑色")
    print("  • 图片处理: 居中对齐，无说明文字，智能URL清理")
    print("  • 分割线: 自动忽略")
    
    return method_2_markdown_docx_custom(md_file, str(output_file))


def install_dependencies_custom():
    """Install required dependencies including PIL for image processing"""
    import subprocess
    
    packages = [
        'markdown',
        'python-docx',
        'requests',
        'Pillow',  # 添加PIL用于图片处理
    ]
    
    print("📦 Installing dependencies...")
    for package in packages:
        try:
            subprocess.check_call([sys.executable, '-m', 'pip', 'install', package])
            print(f"✅ Installed: {package}")
        except subprocess.CalledProcessError:
            print(f"❌ Failed to install: {package}")


def main():
    """Main function with command line interface"""
    parser = argparse.ArgumentParser(description='Convert Markdown files to DOCX format (Custom Format)')
    parser.add_argument('input', nargs='?', help='Input markdown file')
    parser.add_argument('-o', '--output', help='Output DOCX file (optional)')
    parser.add_argument('--install-deps', action='store_true',
                       help='Install required dependencies')
    parser.add_argument('--debug', action='store_true',
                       help='Enable debug output')
    
    args = parser.parse_args()
    
    if args.install_deps:
        install_dependencies_custom()
        return
    
    if not args.input:
        print("🚀 Markdown to DOCX Converter (Custom Format)")
        print("=" * 55)
        print("\n✨ 自定义格式特性:")
        print("  • 一级标题 (#): 保持原样")
        print("  • 二级标题 (##): 宋体 20pt 蓝色 (无编号)")
        print("  • 三级标题 (###): 宋体 18pt 蓝色 + 自动编号 (每个二级标题下从01开始)")
        print("  • 四级标题 (####): 宋体 16pt 蓝色")
        print("  • 五级及以下标题: 宋体 15pt 蓝色")
        print("  • 正文段落: 宋体 15pt 黑色")
        print("  • 列表项: 宋体 15pt 黑色")
        print("  • 代码块: Consolas 10pt")
        print("\n🎨 颜色设置:")
        print("  • 标题颜色: 深蓝色 (RGB: 0, 112, 192)")
        print("  • 正文颜色: 黑色")
        print("\n🔢 自动编号:")
        print("  • 三级标题会自动添加 01 02 03... 编号前缀")
        print("  • 每个二级标题下的三级标题编号从01重新开始")
        print("  • 二级标题不添加编号")
        print("\n📷 图片处理:")
        print("  • 支持本地文件和网络URL")
        print("  • 自动清理URL参数和CDN后缀")
        print("  • 智能路径修复和文件复制")
        print("  • 自动居中对齐")
        print("  • 智能尺寸调整")
        print("  • 不显示图片说明文字")
        print("\n✂️ 其他特性:")
        print("  • 自动忽略markdown分割线 (---, ___, ***)")
        print("\nUsage examples:")
        print("  python md_to_docx_custom.py document.md")
        print("  python md_to_docx_custom.py document.md -o output.docx")
        print("  python md_to_docx_custom.py --install-deps")
        
        # Interactive mode
        md_file = input("\n📄 Enter markdown file path: ").strip()
        if md_file and os.path.exists(md_file):
            success = convert_md_to_docx_custom(md_file)
            if success:
                print("✅ Conversion completed successfully!")
            else:
                print("❌ Conversion failed. Check the output above for details.")
        elif md_file:
            print(f"❌ File not found: {md_file}")
        return
    
    success = convert_md_to_docx_custom(args.input, args.output)
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()